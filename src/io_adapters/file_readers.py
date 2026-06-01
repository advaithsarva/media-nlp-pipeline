import os
import json
import hashlib
from pathlib import Path
from typing import Dict, Any, List, Optional
from datetime import datetime
from abc import ABC, abstractmethod

# Core libraries
import charset_normalizer
import magic

# File-specific processors
import pdfplumber
import fitz  # PyMuPDF
from docx import Document
from bs4 import BeautifulSoup
import markdown as md
import pandas as pd
import csv
from PIL import Image
import pytesseract

# Database
import sqlalchemy
from pymongo import MongoClient

# Compression
import zipfile
import tarfile
import gzip


class BaseReader(ABC):
    """Base class for all file readers"""
    
    def __init__(self, output_dir: str = "./jsonl_output"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.mime_detector = magic.Magic(mime=True)
    
    def detect_encoding(self, file_path: str) -> str:
        """Detect file encoding with fallback"""
        with open(file_path, 'rb') as f:
            raw_data = f.read(100000)
        
        try:
            result = charset_normalizer.from_bytes(raw_data).best()
            if result:
                return result.encoding
        except:
            pass
        
        # Fallback strategy
        for encoding in ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']:
            try:
                raw_data.decode(encoding)
                return encoding
            except:
                continue
        
        return 'utf-8'
    
    def calculate_hash(self, file_path: str) -> str:
        """Calculate SHA256 hash"""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()
    
    def get_base_metadata(self, file_path: str) -> Dict[str, Any]:
        """Get common metadata for all files"""
        path = Path(file_path)
        stat = path.stat()
        
        return {
            'file_path': str(path.absolute()),
            'file_name': path.name,
            'file_size_bytes': stat.st_size,
            'file_extension': path.suffix.lower(),
            'mime_type': self.mime_detector.from_file(str(path)),
            'file_hash_sha256': self.calculate_hash(file_path),
            'created_timestamp': datetime.fromtimestamp(stat.st_ctime).isoformat(),
            'modified_timestamp': datetime.fromtimestamp(stat.st_mtime).isoformat(),
            'processed_timestamp': datetime.now().isoformat(),
            'reader_class': self.__class__.__name__
        }
    
    def write_to_jsonl(self, data: Dict[str, Any], output_file: str):
        """Write single record to JSONL"""
        output_path = self.output_dir / output_file
        
        with open(output_path, 'a', encoding='utf-8') as f:
            f.write(json.dumps(data, ensure_ascii=False) + '\n')
    
    @abstractmethod
    def read(self, file_path: str) -> Dict[str, Any]:
        """Abstract method - must be implemented by subclasses"""
        pass
    
    def process_and_save(self, file_path: str, output_file: str = "output.jsonl"):
        """Process file and save to JSONL"""
        try:
            data = self.read(file_path)
            self.write_to_jsonl(data, output_file)
            return True
        except Exception as e:
            error_data = {
                'file_path': file_path,
                'error': str(e),
                'timestamp': datetime.now().isoformat(),
                'reader_class': self.__class__.__name__
            }
            self.write_to_jsonl(error_data, "errors.jsonl")
            return False
    def _extract_strings_only(self, obj):
        """
        Recursively keep only string values.
        Preserves structure.
        """
        if isinstance(obj, str):
            return obj

        if isinstance(obj, dict):
            filtered = {}
            for k, v in obj.items():
                result = self._extract_strings_only(v)
                if result is not None:
                    filtered[k] = result
            return filtered if filtered else None

        if isinstance(obj, list):
            filtered_list = []
            for item in obj:
                result = self._extract_strings_only(item)
                if result is not None:
                    filtered_list.append(result)
            return filtered_list if filtered_list else None

        # drop everything else (int, float, bool, null, etc.)
        return None


class TxtReader(BaseReader):
    """Reader for .txt, .text files"""
    
    def read(self, file_path: str) -> Dict[str, Any]:
        metadata = self.get_base_metadata(file_path)
        encoding = self.detect_encoding(file_path)
        
        with open(file_path, 'r', encoding=encoding, errors='replace') as f:
            content = f.read()
        
        # Extract RAW features (no preprocessing)
        lines = content.split('\n')
        
        return {
            **metadata,
            'encoding_detected': encoding,
            'raw_content': content,
            'raw_lines': lines,
            'line_count': len(lines),
            'char_count': len(content),
            'byte_count': len(content.encode(encoding, errors='replace')),
            'empty_lines': sum(1 for line in lines if not line.strip()),
            'max_line_length': max(len(line) for line in lines) if lines else 0,
            'contains_null_bytes': '\x00' in content,
            'contains_bom': content.startswith('\ufeff'),
            'line_endings': self._detect_line_endings(file_path),
            'whitespace_stats': {
                'spaces': content.count(' '),
                'tabs': content.count('\t'),
                'newlines': content.count('\n')
            }
        }
    
    def _detect_line_endings(self, file_path: str) -> str:
        """Detect line ending type"""
        with open(file_path, 'rb') as f:
            sample = f.read(1000)
        
        if b'\r\n' in sample:
            return 'CRLF (Windows)'
        elif b'\n' in sample:
            return 'LF (Unix)'
        elif b'\r' in sample:
            return 'CR (Mac)'
        return 'Unknown'


class MarkdownReader(BaseReader):
    """Reader for .md, .markdown files"""
    
    def read(self, file_path: str) -> Dict[str, Any]:
        metadata = self.get_base_metadata(file_path)
        encoding = self.detect_encoding(file_path)
        
        with open(file_path, 'r', encoding=encoding, errors='replace') as f:
            raw_markdown = f.read()
        
        # Convert to HTML for structure analysis
        html_content = md.markdown(raw_markdown, extensions=['extra', 'codehilite', 'toc'])
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Extract structure
        headings = []
        for tag in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
            headings.append({
                'level': int(tag.name[1]),
                'text': tag.get_text(strip=True)
            })
        
        # Extract code blocks
        code_blocks = []
        for pre in soup.find_all('pre'):
            code = pre.find('code')
            if code:
                code_blocks.append({
                    'language': code.get('class', [''])[0].replace('language-', ''),
                    'content': code.get_text()
                })
        
        # Extract links
        links = [{'text': a.get_text(), 'url': a.get('href')} 
                 for a in soup.find_all('a')]
        
        return {
            **metadata,
            'encoding_detected': encoding,
            'raw_markdown': raw_markdown,
            'rendered_html': html_content,
            'plain_text': soup.get_text(separator='\n'),
            'structure': {
                'headings': headings,
                'heading_count': len(headings),
                'code_blocks': code_blocks,
                'code_block_count': len(code_blocks),
                'links': links,
                'link_count': len(links),
                'list_count': len(soup.find_all(['ul', 'ol'])),
                'table_count': len(soup.find_all('table'))
            },
            'markdown_features': {
                'has_frontmatter': raw_markdown.startswith('---'),
                'has_code_fence': '```' in raw_markdown,
                'has_inline_code': '`' in raw_markdown,
                'has_blockquote': raw_markdown.count('> ') > 0,
                'has_images': '![' in raw_markdown
            }
        }


class PDFReader(BaseReader):
    """Reader for .pdf files with comprehensive extraction"""
    
    def read(self, file_path: str) -> Dict[str, Any]:
        metadata = self.get_base_metadata(file_path)
        
        # Primary extraction with pdfplumber
        text_content = []
        tables = []
        page_data = []
        images_info = []
        
        try:
            with pdfplumber.open(file_path) as pdf:
                pdf_metadata = pdf.metadata or {}
                
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text
                    page_text = page.extract_text() or ""
                    
                    # Extract tables
                    page_tables = page.extract_tables() or []
                    
                    # Page dimensions
                    page_info = {
                        'page_number': page_num,
                        'width': page.width,
                        'height': page.height,
                        'rotation': page.rotation,
                        'text_length': len(page_text),
                        'has_text': bool(page_text.strip()),
                        'table_count': len(page_tables)
                    }
                    
                    text_content.append(page_text)
                    tables.extend([{'page': page_num, 'data': table} for table in page_tables])
                    page_data.append(page_info)
                    
                    # Check for images
                    if page.images:
                        images_info.extend([
                            {'page': page_num, 'image_index': i, **img}
                            for i, img in enumerate(page.images)
                        ])
        
        except Exception as e:
            pdf_metadata = {}
            text_content = [f"pdfplumber extraction failed: {e}"]
        
        # Fallback to PyMuPDF if text is empty
        if not ''.join(text_content).strip():
            try:
                doc = fitz.open(file_path)
                text_content = [page.get_text() for page in doc]
                doc.close()
            except:
                pass
        
        full_text = '\n\n'.join(text_content)
        
        # OCR detection (if text is suspiciously short for page count)
        needs_ocr = len(full_text.strip()) < len(page_data) * 50
        
        return {
            **metadata,
            'pdf_metadata': pdf_metadata,
            'page_count': len(page_data),
            'pages': page_data,
            'raw_text_by_page': text_content,
            'full_raw_text': full_text,
            'tables': tables,
            'table_count': len(tables),
            'images': images_info,
            'image_count': len(images_info),
            'text_statistics': {
                'total_chars': len(full_text),
                'total_words': len(full_text.split()),
                'total_lines': len(full_text.split('\n')),
                'avg_chars_per_page': len(full_text) / len(page_data) if page_data else 0
            },
            'quality_indicators': {
                'needs_ocr': needs_ocr,
                'is_scanned': needs_ocr,
                'has_tables': len(tables) > 0,
                'has_images': len(images_info) > 0,
                'is_empty': not full_text.strip()
            },
            'extraction_method': 'pdfplumber_primary_pymupdf_fallback'
        }


class DocsReader(BaseReader):
    """Reader for .docx, .doc files"""
    
    def read(self, file_path: str) -> Dict[str, Any]:
        metadata = self.get_base_metadata(file_path)
        
        # Only .docx is supported directly
        if not file_path.endswith('.docx'):
            return {
                **metadata,
                'error': 'Only .docx format supported. Convert .doc to .docx first.',
                'supported_format': False
            }
        
        doc = Document(file_path)
        
        # Extract paragraphs with styles
        paragraphs = []
        for para in doc.paragraphs:
            paragraphs.append({
                'text': para.text,
                'style': para.style.name,
                'alignment': str(para.alignment) if para.alignment else None,
                'is_heading': para.style.name.startswith('Heading'),
                'runs': len(para.runs)
            })
        
        # Extract tables
        tables = []
        for table_idx, table in enumerate(doc.tables):
            table_data = {
                'table_index': table_idx,
                'rows': len(table.rows),
                'cols': len(table.columns),
                'data': []
            }
            
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data['data'].append(row_data)
            
            tables.append(table_data)
        
        # Extract document properties
        core_props = doc.core_properties
        
        # Get full text
        full_text = '\n\n'.join([p['text'] for p in paragraphs if p['text']])
        
        # Extract headings structure
        headings = [p for p in paragraphs if p['is_heading']]
        
        return {
            **metadata,
            'document_properties': {
                'title': core_props.title,
                'author': core_props.author,
                'subject': core_props.subject,
                'keywords': core_props.keywords,
                'created': core_props.created.isoformat() if core_props.created else None,
                'modified': core_props.modified.isoformat() if core_props.modified else None,
                'revision': core_props.revision
            },
            'paragraphs': paragraphs,
            'paragraph_count': len(paragraphs),
            'full_raw_text': full_text,
            'headings': headings,
            'heading_count': len(headings),
            'tables': tables,
            'table_count': len(tables),
            'sections': len(doc.sections),
            'text_statistics': {
                'total_chars': len(full_text),
                'total_words': len(full_text.split()),
                'total_lines': len(full_text.split('\n'))
            },
            'style_distribution': self._get_style_distribution(paragraphs)
        }
    
    def _get_style_distribution(self, paragraphs: List[Dict]) -> Dict[str, int]:
        """Count paragraph styles"""
        styles = {}
        for para in paragraphs:
            style = para['style']
            styles[style] = styles.get(style, 0) + 1
        return styles



class CSVReader(BaseReader):
    """Reader for .csv, .tsv files"""
    
    def read(self, file_path: str) -> Dict[str, Any]:
        metadata = self.get_base_metadata(file_path)
        encoding = self.detect_encoding(file_path)
        
        # Detect delimiter
        with open(file_path, 'r', encoding=encoding) as f:
            sample = f.read(1024)
            sniffer = csv.Sniffer()
            try:
                dialect = sniffer.sniff(sample)
                delimiter = dialect.delimiter
            except:
                delimiter = ',' if file_path.endswith('.csv') else '\t'
        
        # Read with pandas for robust parsing
        try:
            df = pd.read_csv(file_path, encoding=encoding, delimiter=delimiter)
            
            # Get column info
            columns_info = []
            for col in df.columns:
                col_info = {
                    'name': col,
                    'dtype': str(df[col].dtype),
                    'null_count': int(df[col].isnull().sum()),
                    'unique_count': int(df[col].nunique()),
                    'sample_values': df[col].head(5).tolist()
                }
                columns_info.append(col_info)
            
            # Convert to records (all rows)
            records = df.to_dict('records')
            
            return {
                **metadata,
                'encoding_detected': encoding,
                'delimiter': delimiter,
                'row_count': len(df),
                'column_count': len(df.columns),
                'columns': df.columns.tolist(),
                'columns_info': columns_info,
                'data_types': df.dtypes.astype(str).to_dict(),
                'raw_data': records,
                'sample_data': records[:10],
                'statistics': {
                    'total_cells': df.size,
                    'null_cells': int(df.isnull().sum().sum()),
                    'memory_usage_bytes': int(df.memory_usage(deep=True).sum())
                },
                'data_quality': {
                    'has_null_values': df.isnull().any().any(),
                    'has_duplicates': df.duplicated().any(),
                    'duplicate_count': int(df.duplicated().sum())
                }
            }
        
        except Exception as e:
            # Fallback to basic CSV reading
            with open(file_path, 'r', encoding=encoding) as f:
                reader = csv.reader(f, delimiter=delimiter)
                rows = list(reader)
            
            return {
                **metadata,
                'encoding_detected': encoding,
                'delimiter': delimiter,
                'row_count': len(rows),
                'raw_data': rows,
                'error': f'Pandas parsing failed: {e}'
            }


class JSONReader(BaseReader):
    """Reader for .json files"""
    
    def read(self, file_path: str) -> Dict[str, Any]:
        metadata = self.get_base_metadata(file_path)
        encoding = self.detect_encoding(file_path)
        
        with open(file_path, 'r', encoding=encoding) as f:
            raw_content = f.read()
        
        try:
            json_data = json.loads(raw_content)
            string_only_json = self._extract_strings_only(json_data)

            return {
                **metadata,
                'encoding_detected': encoding,
                'raw_json_string': raw_content,
                'parsed_json': string_only_json,
                'json_structure': {
                    'root_type': type(json_data).__name__,
                    'is_array': isinstance(json_data, list),
                    'is_object': isinstance(json_data, dict),
                    'size': len(json_data) if isinstance(json_data, (list, dict)) else 1,
                    'keys': list(json_data.keys()) if isinstance(json_data, dict) else None,
                    'depth': self._calculate_json_depth(json_data)
                },
                'byte_size': len(raw_content.encode(encoding))
            }
        
        except json.JSONDecodeError as e:
            return {
                **metadata,
                'encoding_detected': encoding,
                'raw_json_string': raw_content,
                'parse_error': str(e),
                'error_line': e.lineno,
                'error_column': e.colno
            }
    
    def _calculate_json_depth(self, obj, current_depth=0):
        """Calculate maximum depth of JSON structure"""
        if isinstance(obj, dict):
            if not obj:
                return current_depth
            return max(self._calculate_json_depth(v, current_depth + 1) for v in obj.values())
        elif isinstance(obj, list):
            if not obj:
                return current_depth
            return max(self._calculate_json_depth(item, current_depth + 1) for item in obj)
        else:
            return current_depth




class JSONLReader(BaseReader):
    """Reader for .jsonl, .ndjson files"""
    
    def read(self, file_path: str) -> Dict[str, Any]:
        metadata = self.get_base_metadata(file_path)
        encoding = self.detect_encoding(file_path)
        
        records = []
        errors = []
        
        with open(file_path, 'r', encoding=encoding) as f:
            for line_num, line in enumerate(f, 1):
                line = line.strip()
                if not line:
                    continue
                
                try:
                    record = json.loads(line)
                    string_only = self._extract_strings_only(record)

                    if string_only:
                        records.append({
                            'line_number': line_num,
                            'data': string_only
                        })
                except json.JSONDecodeError as e:
                    errors.append({
                        'line_number': line_num,
                        'error': str(e),
                        'raw_line': line[:100]
                    })
        
        # Analyze schema consistency
        if records:
            first_keys = set(records[0]['data'].keys()) if isinstance(records[0]['data'], dict) else set()
            schema_consistent = all(
                set(r['data'].keys()) == first_keys 
                for r in records 
                if isinstance(r['data'], dict)
            )
        else:
            schema_consistent = True
            first_keys = set()
        
        return {
            **metadata,
            'encoding_detected': encoding,
            'total_records': len(records),
            'all_records': records,
            'sample_records': records[:10],
            'errors': errors,
            'error_count': len(errors),
            'schema_analysis': {
                'is_consistent': schema_consistent,
                'common_keys': list(first_keys),
                'record_types': list(set(type(r['data']).__name__ for r in records))
            }
        }


class HTMLReader(BaseReader):
    """Reader for .html, .htm files"""
    
    def read(self, file_path: str) -> Dict[str, Any]:
        metadata = self.get_base_metadata(file_path)
        encoding = self.detect_encoding(file_path)
        
        with open(file_path, 'r', encoding=encoding, errors='replace') as f:
            raw_html = f.read()
        
        soup = BeautifulSoup(raw_html, 'lxml')
        
        # Extract metadata tags
        meta_tags = {}
        for meta in soup.find_all('meta'):
            name = meta.get('name') or meta.get('property')
            content = meta.get('content')
            if name and content:
                meta_tags[name] = content
        
        # Extract structure
        headings = []
        for tag in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
            headings.append({
                'level': int(tag.name[1]),
                'text': tag.get_text(strip=True),
                'id': tag.get('id'),
                'class': tag.get('class')
            })
        
        # Extract links
        links = [{'text': a.get_text(strip=True), 'href': a.get('href'), 'rel': a.get('rel')} 
                 for a in soup.find_all('a')]
        
        # Extract images
        images = [{'src': img.get('src'), 'alt': img.get('alt'), 'title': img.get('title')} 
                  for img in soup.find_all('img')]
        
        # Extract scripts and styles
        scripts = [{'src': script.get('src'), 'inline': bool(script.string)} 
                   for script in soup.find_all('script')]
        stylesheets = [link.get('href') for link in soup.find_all('link', rel='stylesheet')]
        
        # Remove script and style elements for clean text
        for element in soup(['script', 'style']):
            element.decompose()
        
        # Extract clean text
        text = soup.get_text(separator='\n')
        lines = (line.strip() for line in text.splitlines())
        clean_text = '\n'.join(line for line in lines if line)
        
        return {
            **metadata,
            'encoding_detected': encoding,
            'raw_html': raw_html,
            'title': soup.title.string if soup.title else None,
            'meta_tags': meta_tags,
            'structure': {
                'headings': headings,
                'heading_count': len(headings),
                'links': links,
                'link_count': len(links),
                'images': images,
                'image_count': len(images),
                'tables': len(soup.find_all('table')),
                'forms': len(soup.find_all('form')),
                'lists': len(soup.find_all(['ul', 'ol']))
            },
            'scripts': scripts,
            'script_count': len(scripts),
            'stylesheets': stylesheets,
            'stylesheet_count': len(stylesheets),
            'extracted_text': clean_text,
            'html_features': {
                'has_doctype': raw_html.strip().lower().startswith('<!doctype'),
                'html_version': self._detect_html_version(raw_html),
                'has_semantic_tags': bool(soup.find_all(['article', 'section', 'nav', 'aside']))
            }
        }
    
    def _detect_html_version(self, html: str) -> str:
        """Detect HTML version from doctype"""
        html_lower = html.lower()
        if '<!doctype html>' in html_lower:
            return 'HTML5'
        elif 'xhtml' in html_lower:
            return 'XHTML'
        elif 'html 4' in html_lower:
            return 'HTML4'
        return 'Unknown'


class XMLReader(BaseReader):
    """Reader for .xml files"""
    
    def read(self, file_path: str) -> Dict[str, Any]:
        metadata = self.get_base_metadata(file_path)
        encoding = self.detect_encoding(file_path)
        
        with open(file_path, 'r', encoding=encoding, errors='replace') as f:
            raw_xml = f.read()
        
        soup = BeautifulSoup(raw_xml, 'lxml-xml')
        
        # Extract root element
        root = soup.find()
        root_info = {
            'name': root.name if root else None,
            'attributes': dict(root.attrs) if root else {},
            'namespace': root.namespace if root and hasattr(root, 'namespace') else None
        }
        
        # Count elements
        all_tags = soup.find_all()
        tag_counts = {}
        for tag in all_tags:
            tag_counts[tag.name] = tag_counts.get(tag.name, 0) + 1
        
        # Extract text content
        text_content = soup.get_text(separator='\n', strip=True)
        
        return {
            **metadata,
            'encoding_detected': encoding,
            'raw_xml': raw_xml,
            'root_element': root_info,
            'element_count': len(all_tags),
            'unique_elements': list(tag_counts.keys()),
            'element_distribution': tag_counts,
            'extracted_text': text_content,
            'xml_features': {
                'has_declaration': raw_xml.strip().startswith('<?xml'),
                'is_well_formed': bool(root),
                'max_depth': self._calculate_xml_depth(soup)
            }
        }
    
    def _calculate_xml_depth(self, element, current_depth=0):
        """Calculate maximum depth of XML tree"""
        children = list(element.children)
        if not children:
            return current_depth
        return max(self._calculate_xml_depth(child, current_depth + 1) 
                   for child in children if hasattr(child, 'children'))


class SQLReader(BaseReader):
    """Reader for SQL databases"""
    
    def read(self, connection_string: str, query: str = None, table_name: str = None) -> Dict[str, Any]:
        """
        Read from SQL database
        connection_string: e.g., 'postgresql://user:pass@localhost/db'
        Either provide query or table_name
        """
        metadata = {
            'connection_string': connection_string.split('@')[-1],  # Hide credentials
            'timestamp': datetime.now().isoformat(),
            'reader_class': self.__class__.__name__
        }
        
        engine = sqlalchemy.create_engine(connection_string)
        
        # If table_name provided, read entire table
        if table_name and not query:
            query = f"SELECT * FROM {table_name}"
        
        if not query:
            # Get database schema
            inspector = sqlalchemy.inspect(engine)
            tables = inspector.get_table_names()
            
            table_info = []
            for table in tables:
                columns = inspector.get_columns(table)
                table_info.append({
                    'table_name': table,
                    'columns': columns,
                    'column_count': len(columns)
                })
            
            return {
                **metadata,
                'database_type': engine.dialect.name,
                'tables': table_info,
                'table_count': len(tables),
                'message': 'Schema information only. Provide query or table_name to extract data.'
            }
        
        # Execute query
        df = pd.read_sql(query, engine)
        
        # Get column info
        columns_info = []
        for col in df.columns:
            col_info = {
                'name': col,
                'dtype': str(df[col].dtype),
                'null_count': int(df[col].isnull().sum()),
                'unique_count': int(df[col].nunique())
            }
            columns_info.append(col_info)
        
        records = df.to_dict('records')
        
        return
# class AudioReader:
#     def __init__(self):
#         pass
               
# class ParquetReader:
#     def __init__(self):
#         pass

# class HadhoopReader:
#     def __init__(self):
#         pass

# class ArchiveReader:
#     def __init__(self):
#         pass

